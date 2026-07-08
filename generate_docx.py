import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- Generate Model & Data Card ---
doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

heading = doc.add_heading('Model & Data Card', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Datasets Used and Sources', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Dataset Name'
hdr_cells[1].text = 'Source / Provider'
hdr_cells[2].text = 'Description & Use Case'

datasets = [
    ("Mirpur 12 Baseline Ground-Truth Data", "Field Survey (Custom)", "In-situ measurements of air temperature, GPS coordinates, vegetation cover, traffic density, and site photos. Used for baseline calibration."),
    ("Landsat 8/9 Thermal Imagery", "USGS Earth Explorer", "Thermal infrared bands used for extracting Land Surface Temperature (LST) and analyzing spatial heat distribution."),
    ("Sentinel-2 Imagery", "Copernicus / ESA", "High-resolution optical data for land cover classification and calculating the Normalized Difference Vegetation Index (NDVI)."),
    ("MODIS Surface Temperature", "NASA Earth Observation", "Used for broad, temporal temperature trend analysis and historical heatwave pattern detection.")
]
for item in datasets:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]

doc.add_paragraph()

doc.add_heading('2. Pre-trained Models and AI Frameworks', level=2)
table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Model / Algorithm Name'
hdr_cells[1].text = 'Provider / Framework'
hdr_cells[2].text = 'License'
hdr_cells[3].text = 'Purpose / Application'

models = [
    ("Linear & Random Forest Regressors", "Scikit-Learn (Python)", "BSD 3-Clause", "Trained on historical temperatures and spatial features to forecast regional surface temperatures 1, 3, and 5 years ahead."),
    ("XGBoost", "DMLC (XGBoost)", "Apache 2.0", "Advanced predictive modeling to capture nonlinear relationships between urban built-up density and heat retention."),
    ("Spatial Interpolation Models", "QGIS / GDAL", "GPL", "Used for kernel density estimation and mapping continuous UHI heat zones from discrete point data.")
]
for item in models:
    row_cells = table2.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]
    row_cells[3].text = item[3]

doc.add_paragraph()

doc.add_heading('3. Known Limitations', level=2)
ul = doc.add_paragraph(style='List Bullet')
ul.add_run('Weather Dependency:').bold = True
ul.add_run(' Satellite remote sensing (especially Landsat and Sentinel) is heavily dependent on cloud-free days. Cloud cover can introduce missing data.')
ul2 = doc.add_paragraph(style='List Bullet')
ul2.add_run('Spatial vs. Temporal Resolution Trade-off:').bold = True
ul2.add_run(' While Landsat provides good spatial resolution, its revisit time is 16 days, limiting real-time daily heat tracking.')
ul3 = doc.add_paragraph(style='List Bullet')
ul3.add_run('Micro-climate Variations:').bold = True
ul3.add_run(' Ground sensors may capture extreme micro-climate effects (e.g., immediate shading from a vehicle) that do not generalize to the 30m pixel of satellite imagery.')

doc.add_heading('4. Ethical Concerns', level=2)
ul4 = doc.add_paragraph(style='List Bullet')
ul4.add_run('Data Privacy:').bold = True
ul4.add_run(' Geotagged ground verification photos collected during field surveys may inadvertently capture identifiable individuals or private properties. (Mitigation: Blurring faces/license plates).')
ul5 = doc.add_paragraph(style='List Bullet')
ul5.add_run('Socioeconomic Misinterpretation:').bold = True
ul5.add_run(' High-resolution heat maps could negatively influence real estate values or insurance premiums in vulnerable areas, exacerbating socioeconomic inequalities.')
ul6 = doc.add_paragraph(style='List Bullet')
ul6.add_run('Algorithmic Bias:').bold = True
ul6.add_run(' If historical field data is predominantly collected from commercial zones while neglecting informal settlements, the predictive models may underestimate heat risks in marginalized communities.')

doc.save('Model_Data_Card.docx')

# --- Generate Project Report ---
doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

heading = doc.add_heading('Project Report: Urban Heat Island Mapping & Predictive GIS Analytics System', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('SciBlitz AI Challenge 2026 | Track B: Environment & Sustainability')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Problem Statement', level=2)
p = doc.add_paragraph('Rapid urbanization significantly alters natural landscapes, replacing vegetation with heat-absorbing materials such as concrete and asphalt. This transformation gives rise to the Urban Heat Island (UHI) effect, where urban regions experience substantially higher temperatures than their rural surroundings. The UHI effect leads to severe consequences, including reduced thermal comfort, heightened energy consumption for cooling, and escalated health risks such as heatstroke and respiratory issues. Despite the growing severity of this issue in cities like Dhaka, urban planners often lack localized, data-driven tools to identify specific micro-heat hotspots and simulate the cooling impact of proposed green infrastructures.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('2. Proposed Solution', level=2)
p = doc.add_paragraph('We propose a comprehensive Urban Heat Island (UHI) Mapping & Analytics System. This system fuses high-resolution remote sensing satellite data with hyper-local ground-truth sensor measurements to create a robust spatial analytics platform. Serving as an Environmental Decision Support System (EDSS), our solution dynamically scales from micro-level analyses (e.g., Mirpur 12) to macroscopic divisional boundaries across Bangladesh.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph('Key features include:')
doc.add_paragraph('Automated Data Fusion: Combining Land Surface Temperature (LST) and Normalized Difference Vegetation Index (NDVI) from satellites with GPS-tagged ground readings.', style='List Bullet')
doc.add_paragraph('Interactive Spatial Dashboard: Visualizing thermal anomalies, heat hotspots, and climate resilience scores via Leaflet.js heatmaps.', style='List Bullet')
doc.add_paragraph('AI Predictive Analytics: Utilizing Machine Learning regression models to forecast future surface temperatures 1, 3, and 5 years ahead based on urbanization trends.', style='List Bullet')
doc.add_paragraph('Green Infrastructure Simulation: Evaluating hypothetical scenarios, such as increasing vegetation by 20%, to predict localized temperature reductions.', style='List Bullet')

doc.add_heading('3. Methodology', level=2)
p = doc.add_paragraph('The project follows a systematic 8-phase pipeline:')
doc.add_paragraph('1. Field Data Collection: Ground temperature, GPS coordinates, vegetation cover, and surface type data were manually recorded using digital sensors across strategic baseline locations (Mirpur 12).', style='List Number')
doc.add_paragraph('2. Remote Sensing: Satellite imagery from Landsat 8/9 and Sentinel-2 was processed via Google Earth Engine to extract thermal infrared bands (LST) and vegetation indices (NDVI).', style='List Number')
doc.add_paragraph('3. Data Preprocessing & Fusion: The disparate datasets were cleaned, normalized, and merged into a unified pandas dataframe using Python.', style='List Number')
doc.add_paragraph('4. GIS Mapping: Spatial interpolation and kernel density estimation were employed to construct continuous thermal maps over the urban grid.', style='List Number')
doc.add_paragraph('5. Heat Risk Index (HRI) Computation: A composite risk score was calculated by normalizing LST and inverting NDVI to categorize zones from "Very Low" to "Extreme" risk.', style='List Number')
doc.add_paragraph('6. Machine Learning Modeling: AI regressors were trained on the fused historical and spatial data to capture the nonlinear dynamics of urban heat retention.', style='List Number')
doc.add_paragraph('7. Dashboard Deployment: The backend was built using Flask, exposing RESTful APIs, while the frontend utilized HTML/CSS/JS and Chart.js for an interactive user experience.', style='List Number')
doc.add_paragraph('8. Validation: Ground-truth temperatures were correlated with satellite LST readings to ensure model accuracy and physical reliability.', style='List Number')

doc.add_heading('4. AI/ML Approach', level=2)
doc.add_heading('4.1 Temperature Forecasting', level=3)
p = doc.add_paragraph('We deployed a suite of machine learning algorithms, primarily Linear Regression and Random Forest Regressors, available via Scikit-Learn. The models utilize historical climatic logs alongside urban built-up density and NDVI features to predict multi-year regional temperature trends. By establishing a baseline year, the AI extrapolates the expected LST 1, 3, and 5 years into the future.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('4.2 Heatwave Alerts & Resilience Scoring', level=3)
p = doc.add_paragraph('An automated monitoring engine evaluates real-time temperature streams against dynamic thresholds (Warning >35°C, Severe >38°C, Extreme >40°C). Concurrently, the AI formulates a Climate Resilience Score (0-100) based on a weighted composite of NDVI (25%), inverted average temperature (25%), inverted Heat Risk Index (25%), and green coverage percentage (25%).')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('4.3 Green Simulation Engine', level=3)
p = doc.add_paragraph('A rule-based simulation engine leverages the trained regression parameters (alpha intercept and beta slope) to calculate the marginal cooling effect of increased vegetation. For instance, simulating a 15% increase in NDVI allows urban planners to instantly visualize the estimated temperature reduction, aiding in targeted afforestation efforts.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading('5. Results', level=2)
p = doc.add_paragraph('The deployment of the system yielded highly actionable insights:')
doc.add_paragraph('Hotspot Identification: The system successfully pinpointed peak heat zones, such as the Mirpur-10 Roundabout (43.9°C), juxtaposed against cool spots like the National Botanical Garden.', style='List Bullet')
doc.add_paragraph('Model Accuracy: The predictive temperature models achieved robust evaluation metrics, maintaining a low Root Mean Square Error (RMSE) and High R-squared score during cross-validation.', style='List Bullet')
doc.add_paragraph('Actionable Recommendations: The EDSS successfully generated tiered intervention plans. High-risk zones triggered aggressive recommendations including vertical gardens, cool pavements, and green roofs, seamlessly exported via automated PDF reports.', style='List Bullet')

doc.add_heading('6. Limitations & Future Work', level=2)
doc.add_heading('6.1 Limitations', level=3)
doc.add_paragraph('The optical and thermal satellite data collection is highly susceptible to cloud cover, limiting continuous temporal monitoring during monsoon seasons.', style='List Bullet')
doc.add_paragraph('Ground-truth data collection requires significant manual effort and is constrained by spatial coverage limitations.', style='List Bullet')
doc.add_paragraph('The predictive models currently assume a linear correlation between NDVI growth and cooling, potentially oversimplifying complex micro-meteorological variables like wind patterns.', style='List Bullet')

doc.add_heading('6.2 Future Work', level=3)
doc.add_paragraph('Real-Time IoT Integration: Deploying stationary, low-cost DHT22/ESP32 sensor nodes across the city to feed live temperature telemetry directly into the ML pipeline.', style='List Bullet')
doc.add_paragraph('Mobile Crowdsourcing App: Developing a public mobile application to allow citizens to report extreme heat zones, crowdsourcing ground-truth data on a massive scale.', style='List Bullet')
doc.add_paragraph('Advanced Deep Learning: Transitioning from Random Forest to Convolutional Neural Networks (CNNs) designed for spatial-temporal satellite image forecasting to capture neighborhood-level thermal spillover effects.', style='List Bullet')

doc.save('Project_Report.docx')
